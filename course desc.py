from docx import Document  # for everything in the document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # to align test left to right
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx2pdf import convert
import os
import requests


def set_indent(myObj):
    myObj.paragraph_format.left_indent = Inches(-0.75)
    myObj.paragraph_format.right_indent = Inches(-0.75)


def first_text(document_element):
    p1 = document_element.add_paragraph("Detailed Syllabus")
    p1.alignment = 1
    set_indent(p1)
    p2 = document_element.add_paragraph('')
    p2.alignment = 1
    set_indent(p2)
    p2.add_run("Lecture-wise Breakup").bold = True


def bold_text(table1, row, col, text):
    x = table1.cell(row, col).paragraphs[0]
    x.alignment = 1
    x.add_run(text).bold = 1
    table1.cell(row, col).vertical_alignment = 1
    table1.cell(row, col).alignment = WD_ALIGN_PARAGRAPH.CENTER


def normal_text(table1, row, col, text):
    x = table1.cell(row, col).paragraphs[0]
    x.alignment = 1
    x.add_run(str(text))
    table1.cell(row, col).vertical_alignment = 1
    table1.cell(row, col).alignment = WD_ALIGN_PARAGRAPH.CENTER


def General_info(document_element, x):
    table1 = document_element.add_table(rows=3, cols=4, style='Table Grid')
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    table1.allow_autofit = False
    for row in table1.rows:
        row.height = Inches(0.5)

    # row 1
    # cell 00
    bold_text(table1, 0, 0, 'Subject Code')
    # cell 01
    normal_text(table1, 0, 1, x[0])
    # cell 02
    bold_text(table1, 0, 2, 'Semester: ' + x[1])
    # cell 03
    table1.cell(0, 3).text = 'Semester 6th Session Month from Feb-June 2022'

    # row 2
    table1.cell(1, 1).merge(table1.cell(1, 2))
    table1.cell(1, 1).merge(table1.cell(1, 3))

    # cell 10
    bold_text(table1, 1, 0, 'Subject\nName')
    # cell 11
    normal_text(table1, 1, 1, x[2])

    # row 3
    # cell 30
    bold_text(table1, 2, 0, 'Credits')
    # cell 31
    normal_text(table1, 2, 1, x[3])
    # cell 32
    bold_text(table1, 2, 2, 'Contact Hours')
    # cell 33
    normal_text(table1, 2, 3, '3-0-0')


def Teachers(document_element, ft):
    document_element.add_paragraph()
    table = document_element.add_table(rows=2, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(1, 2).merge(table.cell(1, 3))
    for row in table.rows:
        row.height = Inches(0.7)

    coordinator_name = "1. {}".format(ft[0]['faculty_name'])
    teachername = ""
    for i in range(1, len(ft)):
        q = ft[i]['faculty_name']
        teachername = teachername + "{}. {} ".format(str(i), q)

    # 00 and 10
    bold_text(table, 0, 0, 'Faculty Names')

    # 01
    bold_text(table, 0, 1, 'Coordinator(s)')

    # 11
    bold_text(table, 1, 1, 'Teacher(s)')

    # 02 and 03
    normal_text(table, 0, 2, coordinator_name)

    # 12 and 13
    normal_text(table, 1, 2, teachername)


def Course_Outcomes(document_element, course_outcomes_data):
    document_element.add_paragraph()
    table = document_element.add_table(rows=len(course_outcomes_data) + 1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    i = 1
    table.cell(0, 0).merge(table.cell(0, 1))

    # 00 and 01
    bold_text(table, 0, 0, 'COURSE OUTCOMES')

    # 02
    bold_text(table, 0, 2, 'COGNITIVE LEVELS')
    table.rows[0].height = Inches(0.6)

    i = 0
    while i < len(course_outcomes_data):
        bold_text(table, i + 1, 0, course_outcomes_data[i][0])
        normal_text(table, i + 1, 1, course_outcomes_data[i][1])
        normal_text(table, i + 1, 2, course_outcomes_data[i][2])
        i = i + 1
    for i in range(1, len(course_outcomes_data) + 1):
        table.cell(i, 0).width = Inches(1.0)

    for row in table.rows:
        row.height = Inches(0.7)


def modules(document_element, modules_data):
    document_element.add_paragraph()
    table = document_element.add_table(rows=len(modules_data) + 2, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    i = 1
    # 00
    bold_text(table, 0, 0, 'Module No.')

    # 01
    bold_text(table, 0, 1, 'Subtitle of Module')

    # 02
    bold_text(table, 0, 2, 'Topics in the Module')

    # 03
    bold_text(table, 0, 3, 'No. Of lectures for the module')

    table.rows[0].height = Inches(0.6)

    i = 0
    total_lec = 0
    while i < len(modules_data):
        normal_text(table, i + 1, 0, modules_data[i][0])
        normal_text(table, i + 1, 1, modules_data[i][1])
        normal_text(table, i + 1, 2, modules_data[i][2])
        normal_text(table, i + 1, 3, str(modules_data[i][3]))
        total_lec = total_lec + modules_data[i][3]
        i = i + 1

    table.cell(len(modules_data) + 1, 0).merge(table.cell(len(modules_data) + 1, 1))
    table.cell(len(modules_data) + 1, 0).merge(table.cell(len(modules_data) + 1, 2))
    bold_text(table, len(modules_data) + 1, 0, 'Total number of Lectures')
    bold_text(table, len(modules_data) + 1, 3, str(total_lec))

    for i in range(0, len(modules_data) + 2):
        table.cell(i, 0).width = Inches(0.75)
        table.cell(i, 1).width = Inches(1.1)
        table.cell(i, 2).width = Inches(4.25)
        table.cell(i, 3).width = Inches(1)


def Evaluation_Criteria(document_element):
    document_element.add_paragraph()
    table = document_element.add_table(rows=7, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    table.cell(0, 0).merge(table.cell(0, 1))
    bold_text(table, 0, 0, 'Evaluation Criteria')
    bold_text(table, 1, 0, 'Components')
    bold_text(table, 1, 1, 'Maximum Marks')
    normal_text(table, 2, 0, 'T1')
    normal_text(table, 2, 1, '20')
    normal_text(table, 3, 0, 'T2')
    normal_text(table, 3, 1, '20')
    normal_text(table, 4, 0, 'End semester examination')
    normal_text(table, 4, 1, '35')
    normal_text(table, 5, 0, 'TA')
    normal_text(table, 5, 1, '25\nAttendance=5, Assignments/Mini-Project/Tutorial/Quiz=20')
    normal_text(table, 6, 0, 'Total')
    normal_text(table, 6, 1, '100')


def Books(document_element, books_data):
    document_element.add_paragraph()
    table = document_element.add_table(rows=len(books_data) + 1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    table.cell(0, 0).merge(table.cell(0, 1))
    p = table.cell(0, 0).paragraphs[0]
    p.add_run('Reading Material: ').bold = True
    p.add_run(
        'Author(s), Title, Edition, Publisher, Year of Publication etc. ( Text books,Reference Books, Journals, '
        'Reports, Websites etc. in the IEEE format)')

    for i in range(1, len(books_data) + 1):
        normal_text(table, i, 0, str(i))
        normal_text(table, i, 1, books_data[i - 1])

    for i in range(1, len(books_data) + 1):
        table.cell(i, 0).width = Inches(0.5)
        table.cell(i, 1).width = Inches(6)


# core object
URL = "http://localhost:5000/api/getCourseDescription/6279265412431d82277b0327"
response = requests.get(URL)
if response.status_code != 200:
    print("Error")
    exit()
response_data = response.json()
if response_data is None:
    print('None type, No data recived, please try back after some time')
y = response_data['message']
module_table = y['Module_table']
text_book_table = y['text_Book_table']
semester = y['semester']
course_ref = y['course_ref']
faculty_table = course_ref['faculty_table']
course_code = course_ref['course_code']
course_name = course_ref['course_name']
course_credit = course_ref['course_credits']
branch = course_ref['Branch']
course_outcomes = course_ref['course_outcome']
reference_book_table = y['reference_books_table']

dc = Document()
sections = dc.sections
section = sections[0]
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
first_text(dc)
General_info(dc, [course_code, semester, course_name, course_credit])
Teachers(dc, faculty_table)
co_final = []
for i in course_outcomes:
    temp = [i['index'], i['Description'], i['cognitive_level']]
    co_final.append(temp)

Course_Outcomes(dc, co_final)
# dc.save("Course Description.docx")

# print(module_table)
modules_data = [
    ['1.', 'Introduction', 'Network terminologies, Network Models, Protocol layers and their services, Connection '
                           'Oriented and Connectionless services, Physical Media. ', 4],
    ['2.', 'The Application Layer', 'Principles of Application-Layer Protocols, HTTP, File Transfer: FTP, DNS, '
                                    'Electronic Mail in the Internet', 4],
    ['3.', 'The Transport Layer', 'Transport-Layer Services and Principles, Multiplexing and Demultiplexing '
                                  'Applications, UDP and TCP, Connection Establishment, Transport Layer Protocols '
                                  '(go back N, stop and wait, selective repeat), Flow Control,  TCP Congestion '
                                  'Control', 8]]
modules(dc, modules_data)
Evaluation_Criteria(dc)
print(text_book_table)
books=[]
for i in text_book_table:
    s = "{} by {}".format(i['Name'], i['Author'])
    books.append(s)
Books(dc, books)
dc.save("Course Description.docx")
'''
dc.add_page_break()
p = dc.add_paragraph()
p.add_run('Project based learning: ').bold = True
p.add_run(
    'Each student in a group of 2-4 will choose some real-world problems such as congestion control, building smart '
    'devices, network traffic analyser etc. for development and analysis. By applying the different network protocol '
    'layer concepts and with the help of simulators it helps the students in enhancing their understanding and skills '
    'towards networking, communication and IoT related issues leading towards employability in IT and hardware '
    'sector.')

'''
a = input("Do you want a word(w) or PDF(p)?")

if a in ['w', 'W']:
    dc.save("Course Description.docx")
elif a in ['p', 'P']:
    dc.save("Course Description.docx")
    convert("Course Description.docx", "Course Description.pdf")
    os.remove("Course Description.docx")
