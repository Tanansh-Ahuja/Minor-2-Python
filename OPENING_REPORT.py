from docx import Document  # for everything in the document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # to align test left to right
# from docx.enum.table import WD_ALIGN_VERTICAL  # to align text top to down
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
# from docx2pdf import convert
from docx.shared import Pt


# import os


def set_indent(a):
    a.paragraph_format.left_indent = Inches(-0.75)
    a.paragraph_format.right_indent = Inches(-0.75)


def bold_text(table1, row, col, text):
    x = table1.cell(row, col).paragraphs[0]
    x.alignment = 1
    x.add_run(text).bold = 1
    table1.cell(row, col).vertical_alignment = 1
    table1.cell(row, col).alignment = WD_ALIGN_PARAGRAPH.CENTER


def bold_text2(table1, row, col, text, size):
    x = table1.cell(row, col).paragraphs[0]
    x.alignment = 1
    y = x.add_run(text)
    y.bold = 1
    y.font.size = Pt(size)
    table1.cell(row, col).vertical_alignment = 1
    table1.cell(row, col).alignment = WD_ALIGN_PARAGRAPH.CENTER


def normal_text(table1, row, col, text):
    x = table1.cell(row, col).paragraphs[0]
    x.alignment = 1
    x.add_run(text)
    table1.cell(row, col).vertical_alignment = 1
    table1.cell(row, col).alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_heading(dc, a, b, c):
    p = dc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(a)
    r.font.size = Pt(b)
    r.bold = True
    p.paragraph_format.space_after = 1
    if 'u' == c:
        r.underline = True


def set_first_block(document_element, a, b):
    for i in a:
        p = document_element.add_paragraph()
        r = p.add_run(i)
        r.font.size = Pt(b)
        r.bold = True
        p.paragraph_format.space_after = 1

        # print(p.paragraph_format.line_spacing)


def set_courseoutcomes(dc, course_outcomes):
    dc.add_paragraph()
    p = dc.add_paragraph()
    r = p.add_run('1. Course Outcomes:')
    r.bold = 1
    r.underline = 1
    r.font.size = Pt(12)
    dc.add_paragraph('At the completion of the course, students will be able to, ')
    table = dc.add_table(rows=len(course_outcomes) + 1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.cell(0, 0).merge(table.cell(0, 1))

    # 00 and 01
    bold_text(table, 0, 0, 'COURSE OUTCOMES')

    # 02
    bold_text(table, 0, 2, 'COGNITIVE LEVELS')
    table.rows[0].height = Inches(0.6)

    i = 0
    while i < len(course_outcomes):
        bold_text(table, i + 1, 0, course_outcomes[i][0])
        normal_text(table, i + 1, 1, course_outcomes[i][1])
        normal_text(table, i + 1, 2, course_outcomes[i][2])
        i = i + 1
    for i in range(1, len(course_outcomes) + 1):
        table.cell(i, 0).width = Inches(1.0)

    for row in table.rows:
        row.height = Inches(0.7)


def co_po_mapping(document_element, co_po_mappings_data):
    p = document_element.add_paragraph()
    r = p.add_run('2. CO-PO and CO-PSO Mapping:')
    r.bold = 1
    r.underline = 1
    l = len(co_po_mappings_data)
    r.font.size = Pt(12)
    table = document_element.add_table(rows=l + 2, cols=15, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    bold_text2(table, 0, 0, 'CO\'s', 10)
    x = 1
    for i in range(1, 10):
        bold_text2(table, 0, x, 'PO' + str(i), 10)
        x = x + 1
    bold_text2(table, 0, 10, 'PO10', 9.5)
    bold_text2(table, 0, 11, 'PO11', 9.5)
    bold_text2(table, 0, 12, 'PO12', 9.5)
    bold_text2(table, 0, 13, 'PS01', 10)
    bold_text2(table, 0, 14, 'PS02', 10)
    for i in range(0, l):
        for j in range(0, 15):
            normal_text(table, i + 1, j, str(co_po_mappings_data[i][j]))

    bold_text(table, l + 1, 0, 'Avg')
    q = [0 for _ in range(14)]
    for i in range(l):
        for j in range(0, len(co_po_mappings_data[0]) - 1):
            q[j] = q[j] + co_po_mappings_data[i][j + 1]

    for i in range(len(q)):
        q[i] = q[i] / l

    for i in range(0, 14):
        normal_text(table, l + 1, i + 1, str(q[i]))


def identify_gaps(document_element, identify_gaps_data):
    if len(identify_gaps_data) == 0:
        return
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('3. Identified gaps in Syllabus/ Course Description (If Any): ')
    r.bold = 1
    r.underline = 1
    table = document_element.add_table(rows=len(identify_gaps_data) + 1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    bold_text(table, 0, 0, 'Topics to be introduced')
    bold_text(table, 0, 1, 'Strengthens CO')
    bold_text(table, 0, 2, 'Strengthens PO, PSO')
    bold_text(table, 0, 3, 'Method of Identification')

    for i in range(len(identify_gaps_data)):
        normal_text(table, i + 1, 0, identify_gaps_data[i][0])
        normal_text(table, i + 1, 1, identify_gaps_data[i][1])
        normal_text(table, i + 1, 2, identify_gaps_data[i][2])
        normal_text(table, i + 1, 3, identify_gaps_data[i][3])
    for i in range(0, len(identify_gaps_data) + 1):
        table.cell(i, 2).width = Inches(1.5)
        table.cell(i, 3).width = Inches(3.0)
        table.cell(i, 1).width = Inches(1.0)


def modification_in_syllabus(document_element, mod_in_syllabus_data):
    if len(mod_in_syllabus_data) == 0:
        return
    document_element.add_paragraph()
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('4. Modifications in Syllabus/ Course Description (If Any)')
    r.bold = 1
    r.underline = 1
    table = document_element.add_table(rows=len(mod_in_syllabus_data) + 1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    bold_text(table, 0, 0, 'Details of Modification (Addition/ Removal)')
    bold_text(table, 0, 1, 'Justification')
    bold_text(table, 0, 2, 'Strengthens POs/PSOs ')
    for i in range(len(mod_in_syllabus_data)):
        normal_text(table, i + 1, 0, mod_in_syllabus_data[i][0])
        normal_text(table, i + 1, 1, mod_in_syllabus_data[i][1])
        normal_text(table, i + 1, 2, mod_in_syllabus_data[i][2])
    for i in range(0, len(mod_in_syllabus_data) + 1):
        table.cell(i, 0).width = Inches(2)
        table.cell(i, 1).width = Inches(3.5)
        table.cell(i, 2).width = Inches(1.75)


def improve_attainment(document_element, attainment_data):
    if attainment_data is None or len(attainment_data) == 0:
        return
    document_element.add_paragraph()
    document_element.add_paragraph()
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('5. Actions for Improving CO Attainments: ')
    r.bold = 1
    r.underline = 1
    table = document_element.add_table(rows=len(attainment_data) + 1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    bold_text(table, 0, 0, 'CO\'s')
    bold_text(table, 0, 1, 'Attainment')
    bold_text(table, 0, 2, 'Action to be taken to improve CO attainment ')
    bold_text(table, 0, 3, 'Strengthens POs/PSOs ')
    for i in range(len(attainment_data)):
        normal_text(table, i + 1, 0, attainment_data[i][0])
        normal_text(table, i + 1, 1, attainment_data[i][1])
        normal_text(table, i + 1, 2, attainment_data[i][2])
        normal_text(table, i + 1, 3, attainment_data[i][3])
    for i in range(0, len(attainment_data) + 1):
        table.cell(i, 0).width = Inches(1.0)
        table.cell(i, 1).width = Inches(2.0)
        table.cell(i, 2).width = Inches(3.0)
        table.cell(i, 3).width = Inches(2.0)


def innovative_teaching(document_element, s):
    if s is None or len(s) == 0:
        return
    document_element.add_paragraph()
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('6. Innovative Teaching and Learning Method to be used: ')
    r.bold = 1
    r.underline = 1
    q = p.add_run(s)


def strategies(document_element, strategy):
    if strategy is None or len(strategy) == 0:
        return
    document_element.add_paragraph()
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('7. Strategies for:\n')
    r.bold = 1
    r.underline = 1
    p = document_element.add_paragraph()
    r = p.add_run('Weak Learners: ')
    r.bold = 1
    q = p.add_run(strategy[0])
    p = document_element.add_paragraph()
    r = p.add_run('Bright Students: ')
    r.bold = 1
    q = p.add_run(strategy[1])


def innovative_eval_strategy(document_element, s):
    if s is None or len(s) == 0:
        return
    document_element.add_paragraph()
    document_element.add_paragraph()
    p = document_element.add_paragraph()
    r = p.add_run('8. Innovative Evaluation Strategy to be used (If any):\n')
    r.bold = 1
    q = p.add_run(s)


# MAIN

dc = Document()
style = dc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
section = dc.sections[0]
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(0.50)
section.bottom_margin = Inches(0.75)

set_heading(dc, 'Department of Computer Science and Engineering & IT', 16, 'n')
set_heading(dc, 'AY: 2021-22 Odd Semester', 14, 'n')
set_heading(dc, 'Course Opening Report', 14, 'u')
dc.add_paragraph()
set_first_block(dc, ['Programme Name:', 'Semester:', 'Course Name & Code: ', 'Name of Course Coordinator(s):'], 12)

course_outcomes = [
    ['CO1', 'Defining the basics of networking, components and underlying technologies', 'Remembering (Level 1)'],
    ['CO2',
     'Illustrate the various key protocols in OSI model and TCP/IP protocol suite and explain various application '
     'protocols.',
     'Understanding (Level 2)'],
    ['CO3', 'Examine various transport protocols and its performance enhancing mechanisms.', 'Analysing (Level 4)'],
    ['CO4', 'Determine the shortest path for the network using various routing protocols and evaluate it.',
     'Evaluating (Level 5)'],
    ['CO5',
     'Choose IP & MAC addressing mechanisms and data link layer protocols to solve communication, error detection and '
     'correction problems.',
     'Applying (Level 3)'],
    ['CO6',
     'Identification and description of various components, architectures and protocols of Internet of Things (IoT) '
     'and their real life problems.',
     'Understand (Level 2)']
]
dc.add_paragraph()
set_courseoutcomes(dc, course_outcomes)

copo_mappings = [['CO1', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3],
                 ['CO2', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3],
                 ['CO3', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3],
                 ['CO4', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3],
                 ['CO5', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3],
                 ['CO6', 1, 2, 3, 1, 2, 3, 1, 2, 3, 1, 2, 3, 3, 3]]
dc.add_page_break()

co_po_mapping(dc, copo_mappings)

identifygaps = [['topic 1', 'CO1, CO2, CO5', 'PO3,PO6,PO12,PSO12', 'LOREM IPSUM'],
                ['topic 2', 'CO1, CO2, CO5', 'PO3,PO6,PO12,PSO12', 'LOREM IPSUM'],
                ['topic 3', 'CO1, CO2, CO5', 'PO3,PO6,PO12,PSO12', 'LOREM IPSUM']]
dc.add_paragraph()
dc.add_paragraph()
identify_gaps(dc, identifygaps)

mod_syllabus_data = [
    ['Addition', 'some justification', 'PO1, PO3, PO6, PSO1'],
    ['Removing', 'some justification', 'PO1, PO3, PO6, PSO1'],
    ['Addition', 'some justification', 'PO1, PO3, PO6, PSO1']
]

modification_in_syllabus(dc, mod_syllabus_data)

improve_attainment(dc, attainment_data=None)

innovative_teaching(dc, s=None)

strategies(dc, strategy=None)

innovative_eval_strategy(dc, s=None)
dc.add_paragraph()
dc.add_paragraph()
q = dc.add_paragraph()
w = q.add_run('Signature Module Coordinator: ')
w.bold = True
for i in range(2):
    dc.add_paragraph()
q = dc.add_paragraph()
w = q.add_run('Signature Course Coordinator: ')
w.bold = True

dc.save("Opening Report.docx")
